"""Paper writing and compilation tools for the BPS Academic Research Agent.

Provides tools for:
- write_section: Generate or update a LaTeX/Markdown paper section
- compile_paper: Compile LaTeX paper to PDF or DOCX
- generate_table: Create publication-quality academic tables
- generate_diagram: Generate diagrams using Mermaid or TikZ
- convert_figure_tikz: Convert matplotlib figures to TikZ/PGFplots
"""

from __future__ import annotations

import re
import subprocess
import textwrap
from pathlib import Path
from typing import Any

from .base import Tool, ToolResult

# ---------------------------------------------------------------------------
# Section templates for LaTeX
# ---------------------------------------------------------------------------

_LATEX_SECTION_COMMANDS: dict[str, str] = {
    "abstract": r"\begin{abstract}",
    "introduction": r"\section{Introduction}",
    "literature_review": r"\section{Literature Review}",
    "methodology": r"\section{Methodology}",
    "results": r"\section{Results}",
    "discussion": r"\section{Discussion}",
    "conclusion": r"\section{Conclusion}",
}

_LATEX_SECTION_END: dict[str, str] = {
    "abstract": r"\end{abstract}",
}

_SECTION_GUIDELINES: dict[str, str] = {
    "abstract": "150-300 words. Structure: Objective → Methods → Results → Conclusions.",
    "introduction": "Broad context → Current knowledge → Gap → Objectives/hypotheses. Funnel structure.",
    "literature_review": "Thematic organization. Every claim needs \\cite{}. Group by framework, methods, findings.",
    "methodology": "Study design → Population → Instruments → Procedures → Variables → Analysis. Past tense.",
    "results": "Preliminary analyses → Descriptive stats → Main analyses → Robustness checks.",
    "discussion": "Summary → Interpretation vs literature → Implications → Limitations → Future research.",
    "conclusion": "200-400 words. Restate findings. No new data. Practical implications. Future directions.",
}


# ---------------------------------------------------------------------------
# 1. WriteSectionTool
# ---------------------------------------------------------------------------


class WriteSectionTool(Tool):
    """Generate or update a paper section in LaTeX or Markdown format."""

    def __init__(self, workspace_dir: str = "./workspace") -> None:
        self._workspace_dir = workspace_dir

    @property
    def name(self) -> str:
        return "write_section"

    @property
    def description(self) -> str:
        return (
            "Generate or update a paper section. Writes content to "
            "paper/sections/ directory in LaTeX or Markdown format, wraps in "
            "appropriate commands, and updates project.yaml section status."
        )

    @property
    def parameters(self) -> dict[str, Any]:
        return {
            "type": "object",
            "properties": {
                "section_name": {
                    "type": "string",
                    "enum": [
                        "abstract",
                        "introduction",
                        "literature_review",
                        "methodology",
                        "results",
                        "discussion",
                        "conclusion",
                    ],
                    "description": "Name of the paper section to write",
                },
                "content": {
                    "type": "string",
                    "description": "The section text content",
                },
                "format": {
                    "type": "string",
                    "enum": ["latex", "markdown"],
                    "description": "Output format (default: latex)",
                    "default": "latex",
                },
                "append": {
                    "type": "boolean",
                    "description": "Append to existing section instead of replacing",
                    "default": False,
                },
            },
            "required": ["section_name", "content"],
        }

    async def execute(
        self,
        section_name: str,
        content: str,
        format: str = "latex",
        append: bool = False,
        **kwargs: Any,
    ) -> ToolResult:
        """Write a paper section to the workspace."""
        try:
            from mini_agent.research.writing.section_writer import SectionWriter

            writer = SectionWriter(self._workspace_dir)

            if append:
                existing = writer.get_section(section_name)
                if existing:
                    # Strip trailing whitespace/end commands and append
                    plain_existing = re.sub(r"\\end\{abstract\}\s*$", "", existing).rstrip()
                    # Remove header comments from existing for clean append
                    body_lines = [
                        line for line in plain_existing.splitlines() if not line.startswith("%") and line.strip()
                    ]
                    # Skip LaTeX commands at the start
                    body_start = 0
                    for i, line in enumerate(body_lines):
                        if line.startswith("\\section") or line.startswith("\\begin") or line.startswith("\\label"):
                            body_start = i + 1
                        else:
                            break
                    existing_body = "\n".join(body_lines[body_start:])
                    content = existing_body + "\n\n" + content

            output_path = writer.write_section(section_name, content, output_format=format)
            word_count = writer.get_word_count(section_name)

            # Update project.yaml if available
            self._update_project_yaml(Path(self._workspace_dir), section_name, word_count)

            return ToolResult(
                success=True,
                content=(
                    f"## Section Written: {section_name.replace('_', ' ').title()}\n\n"
                    f"**File:** {output_path.relative_to(Path(self._workspace_dir))}\n"
                    f"**Format:** {format}\n"
                    f"**Word count:** {word_count}\n"
                    f"**Append:** {append}\n"
                    f"**Guidelines:** {_SECTION_GUIDELINES.get(section_name, 'N/A')}"
                ),
            )
        except Exception as e:
            return ToolResult(success=False, error=f"Write section failed: {e}")

    def _update_project_yaml(self, ws: Path, section_name: str, word_count: int) -> None:
        """Update project.yaml with section status."""
        try:
            import yaml
        except ImportError:
            return

        yaml_path = ws / "project.yaml"
        if not yaml_path.exists():
            return

        try:
            data = yaml.safe_load(yaml_path.read_text(encoding="utf-8")) or {}
            sections = data.setdefault("sections", {})
            section_info = sections.setdefault(section_name, {})
            section_info["status"] = "draft"
            section_info["word_count"] = word_count
            yaml_path.write_text(yaml.dump(data, default_flow_style=False), encoding="utf-8")
        except Exception:
            pass  # Non-critical


# ---------------------------------------------------------------------------
# 2. CompilePaperTool
# ---------------------------------------------------------------------------


class CompilePaperTool(Tool):
    """Compile paper sections into a final document (PDF or DOCX)."""

    def __init__(self, workspace_dir: str = "./workspace") -> None:
        self._workspace_dir = workspace_dir

    @property
    def name(self) -> str:
        return "compile_paper"

    @property
    def description(self) -> str:
        return (
            "Compile all paper sections into a final document. Supports PDF "
            "(via LaTeX) and DOCX (via python-docx). For PDF, tries tectonic "
            "then pdflatex. Returns path to compiled output."
        )

    @property
    def parameters(self) -> dict[str, Any]:
        return {
            "type": "object",
            "properties": {
                "output_format": {
                    "type": "string",
                    "enum": ["pdf", "tex", "docx"],
                    "description": "Output format: pdf (compiled), tex (combined .tex file), docx",
                    "default": "pdf",
                },
                "template": {
                    "type": "string",
                    "description": "Journal template (elsevier, ieee, springer, etc.)",
                    "default": "elsevier",
                },
            },
            "required": [],
        }

    async def execute(
        self,
        output_format: str = "pdf",
        template: str = "elsevier",
        **kwargs: Any,
    ) -> ToolResult:
        """Compile the paper."""
        try:
            ws = Path(self._workspace_dir)

            if output_format == "pdf":
                return await self._compile_pdf(ws, template)
            elif output_format == "tex":
                return self._compile_tex(ws, template, **kwargs)
            elif output_format == "docx":
                return self._compile_docx(ws, template)
            else:
                return ToolResult(
                    success=False,
                    error=f"Unsupported output format: {output_format}. Use: pdf, tex, or docx",
                )
        except Exception as e:
            return ToolResult(success=False, error=f"Compilation failed: {e}")

    async def _compile_pdf(self, ws: Path, template: str, **kwargs: Any) -> ToolResult:
        """Compile to PDF using pdflatex on a combined .tex file."""
        import subprocess as _sp

        # First generate the combined .tex file
        tex_result = self._compile_tex(ws, template, **kwargs)
        if not tex_result.success:
            return tex_result

        tex_path = ws / "paper" / "compiled" / "paper.tex"
        if not tex_path.exists():
            return ToolResult(success=False, error="Failed to generate paper.tex")

        # Check for pdflatex
        try:
            r = _sp.run(["which", "pdflatex"], capture_output=True, timeout=5)
            if r.returncode != 0:
                return ToolResult(
                    success=False,
                    error=(
                        "pdflatex not found. Install with: sudo apt-get install texlive-latex-base\n"
                        "Or use output_format='tex' and upload to Overleaf."
                    ),
                )
        except (FileNotFoundError, _sp.TimeoutExpired):
            return ToolResult(success=False, error="Cannot check for pdflatex")

        # Compile with pdflatex (run twice for references)
        compile_dir = tex_path.parent
        for _pass in range(2):
            _sp.run(
                ["pdflatex", "-interaction=nonstopmode", "-output-directory", str(compile_dir), str(tex_path)],
                capture_output=True,
                text=True,
                timeout=120,
                cwd=str(compile_dir),
            )

        pdf_path = compile_dir / "paper.pdf"
        if pdf_path.exists():
            return ToolResult(
                success=True,
                content=(
                    f"## PDF Compilation Successful\n\n"
                    f"**Output:** {pdf_path}\n"
                    f"**Size:** {pdf_path.stat().st_size / 1024:.1f} KB\n"
                    f"**Template:** {template}\n"
                    f"**Pages:** (check PDF)"
                ),
            )
        else:
            # Return the log for debugging
            log_path = compile_dir / "paper.log"
            log_tail = ""
            if log_path.exists():
                log_lines = log_path.read_text(encoding="utf-8", errors="replace").splitlines()
                log_tail = "\n".join(log_lines[-20:])
            return ToolResult(
                success=False,
                error=f"PDF compilation failed. Last log lines:\n{log_tail}",
            )

    def _compile_tex(self, ws: Path, template: str, **kwargs: Any) -> ToolResult:
        """Generate a combined .tex file with all sections inlined."""
        from mini_agent.research.writing.section_writer import SectionWriter

        writer = SectionWriter(ws)
        sections = writer.get_all_sections()

        if not sections:
            return ToolResult(
                success=False,
                error="No sections found. Write sections first with write_section.",
            )

        # Build combined LaTeX document
        title = kwargs.get("title", self._get_title(ws))
        authors = kwargs.get("authors", self._get_authors(ws))
        abstract_text = kwargs.get("abstract", "")
        keywords_list = kwargs.get("keywords", [])

        author_str = " \\and ".join(authors) if isinstance(authors, list) else str(authors)
        keywords_str = ", ".join(keywords_list) if keywords_list else ""

        # Detect available document class
        import subprocess as _sp

        _has_elsarticle = False
        try:
            r = _sp.run(["kpsewhich", "elsarticle.cls"], capture_output=True, timeout=5)
            _has_elsarticle = r.returncode == 0 and r.stdout.strip()
        except (FileNotFoundError, _sp.TimeoutExpired):
            pass

        if _has_elsarticle:
            doc_class = r"\documentclass[preprint,12pt]{elsarticle}"
        else:
            doc_class = r"\documentclass[12pt]{article}"

        lines = [
            doc_class,
            r"\usepackage[utf8]{inputenc}",
            r"\usepackage{amsmath,amssymb}",
            r"\usepackage{graphicx}",
            r"\usepackage{hyperref}",
            r"\usepackage{booktabs}",
            r"\usepackage{natbib}",
            r"\usepackage{lineno}",
            r"\linenumbers",
            "",
            f"\\title{{{title}}}",
            f"\\author{{{author_str}}}",
            "",
            r"\begin{document}",
            r"\maketitle",
            "",
        ]

        if abstract_text:
            lines.extend(
                [
                    r"\begin{abstract}",
                    abstract_text,
                    r"\end{abstract}",
                    "",
                ]
            )

        if keywords_str:
            if _has_elsarticle:
                lines.extend(
                    [
                        r"\begin{keyword}",
                        keywords_str,
                        r"\end{keyword}",
                        "",
                    ]
                )
            else:
                lines.extend(
                    [
                        f"\\noindent\\textbf{{Keywords:}} {keywords_str}",
                        "",
                    ]
                )

        # Inline all sections in order
        section_order = ["introduction", "literature_review", "methodology", "results", "discussion", "conclusion"]
        for sec_name in section_order:
            if sec_name in sections:
                lines.append(f"% === {sec_name.upper()} ===")
                lines.append(sections[sec_name])
                lines.append("")

        # Add any remaining sections not in standard order
        for sec_name, content in sections.items():
            if sec_name not in section_order:
                lines.append(f"% === {sec_name.upper()} ===")
                lines.append(content)
                lines.append("")

        lines.extend(
            [
                r"\bibliographystyle{elsarticle-num}",
                r"\bibliography{references}",
                "",
                r"\end{document}",
            ]
        )

        combined_tex = "\n".join(lines)

        # Write to file
        output_dir = ws / "paper" / "compiled"
        output_dir.mkdir(parents=True, exist_ok=True)
        output_path = output_dir / "paper.tex"
        output_path.write_text(combined_tex, encoding="utf-8")

        return ToolResult(
            success=True,
            content=(
                f"## LaTeX File Generated\n\n"
                f"**Output:** {output_path}\n"
                f"**Sections:** {len(sections)}\n"
                f"**Total size:** {len(combined_tex)} characters\n"
                f"**Word count:** ~{len(combined_tex.split())}\n\n"
                f"You can compile this with `pdflatex paper.tex` or upload to Overleaf."
            ),
        )

    def _compile_docx(self, ws: Path, template: str) -> ToolResult:
        """Compile to DOCX using python-docx."""
        try:
            from docx import Document  # type: ignore[import-untyped]
        except ImportError:
            return ToolResult(
                success=False,
                error="python-docx not installed. Install with: pip install python-docx",
            )

        from mini_agent.research.writing.section_writer import SectionWriter

        writer = SectionWriter(ws)
        sections = writer.get_all_sections()

        if not sections:
            return ToolResult(
                success=False,
                error="No sections found. Write sections first with write_section.",
            )

        doc = Document()
        title = self._get_title(ws)
        doc.add_heading(title, level=0)

        section_order = [
            "abstract",
            "introduction",
            "literature_review",
            "methodology",
            "results",
            "discussion",
            "conclusion",
        ]

        for name in section_order:
            if name in sections:
                content = sections[name]
                # Strip LaTeX commands
                plain = re.sub(r"(?m)^%.*$", "", content)
                plain = re.sub(r"\\[a-zA-Z]+\*?(\{[^}]*\})?", "", plain)
                plain = re.sub(r"[{}$\\]", "", plain)
                plain = re.sub(r"\s+", " ", plain).strip()

                heading = name.replace("_", " ").title()
                doc.add_heading(heading, level=1)
                doc.add_paragraph(plain)

        out_dir = ws / "writing" / "compiled"
        out_dir.mkdir(parents=True, exist_ok=True)
        docx_path = out_dir / "paper.docx"
        doc.save(str(docx_path))

        return ToolResult(
            success=True,
            content=(
                f"## DOCX Compilation Successful\n\n"
                f"**Output:** {docx_path}\n"
                f"**Sections:** {len(sections)}\n"
                f"**Template:** {template}"
            ),
        )

    @staticmethod
    def _get_title(ws: Path) -> str:
        """Extract title from project.yaml or return default."""
        try:
            import yaml

            yaml_path = ws / "project.yaml"
            if yaml_path.exists():
                data = yaml.safe_load(yaml_path.read_text(encoding="utf-8")) or {}
                project = data.get("project", {})
                return project.get("title", "Untitled Paper")
        except Exception:
            pass
        return "Untitled Paper"

    @staticmethod
    def _get_authors(ws: Path) -> list[str]:
        """Extract authors from project.yaml or return default."""
        try:
            import yaml

            yaml_path = ws / "project.yaml"
            if yaml_path.exists():
                data = yaml.safe_load(yaml_path.read_text(encoding="utf-8")) or {}
                project = data.get("project", {})
                authors = project.get("authors", [])
                if isinstance(authors, list) and authors:
                    return authors
        except Exception:
            pass
        return ["Author"]


# ---------------------------------------------------------------------------
# 3. GenerateTableTool
# ---------------------------------------------------------------------------


class GenerateTableTool(Tool):
    """Generate publication-quality academic tables."""

    def __init__(self, workspace_dir: str = "./workspace") -> None:
        self._workspace_dir = workspace_dir

    @property
    def name(self) -> str:
        return "generate_table"

    @property
    def description(self) -> str:
        return (
            "Generate a publication-quality academic table from data. Supports "
            "booktabs style for LaTeX and Markdown output. Handles numeric "
            "formatting and significance stars."
        )

    @property
    def parameters(self) -> dict[str, Any]:
        return {
            "type": "object",
            "properties": {
                "data_path": {
                    "type": "string",
                    "description": "Path to data file (CSV or JSON)",
                },
                "columns": {
                    "type": "array",
                    "items": {"type": "string"},
                    "description": "Columns to include (default: all)",
                },
                "caption": {
                    "type": "string",
                    "description": "Table caption",
                },
                "label": {
                    "type": "string",
                    "description": "LaTeX label (e.g., tab:descriptive)",
                },
                "format": {
                    "type": "string",
                    "enum": ["latex", "markdown"],
                    "description": "Output format (default: latex)",
                    "default": "latex",
                },
            },
            "required": ["caption"],
        }

    async def execute(
        self,
        caption: str,
        data_path: str | None = None,
        columns: list[str] | None = None,
        label: str | None = None,
        format: str = "latex",
        **kwargs: Any,
    ) -> ToolResult:
        """Generate an academic table."""
        try:
            import pandas as pd
        except ImportError:
            return ToolResult(
                success=False,
                error="pandas not installed. Install with: pip install pandas",
            )

        try:
            ws = Path(self._workspace_dir)

            # Load data
            if data_path:
                full_path = ws / data_path if not Path(data_path).is_absolute() else Path(data_path)
                if not full_path.exists():
                    return ToolResult(success=False, error=f"Data file not found: {full_path}")
                if full_path.suffix == ".csv":
                    df = pd.read_csv(str(full_path))
                elif full_path.suffix == ".json":
                    df = pd.read_json(str(full_path))
                elif full_path.suffix in (".xlsx", ".xls"):
                    df = pd.read_excel(str(full_path))
                else:
                    return ToolResult(success=False, error=f"Unsupported format: {full_path.suffix}")
            elif "data" in kwargs and kwargs["data"]:
                df = pd.DataFrame(kwargs["data"])
            else:
                return ToolResult(
                    success=False,
                    error="No data provided. Supply data_path or data parameter.",
                )

            # Filter columns
            if columns:
                available = [c for c in columns if c in df.columns]
                if available:
                    df = df[available]

            # Round numeric columns
            for col in df.select_dtypes(include=["float64", "float32"]).columns:
                df[col] = df[col].round(3)

            # Generate label
            if not label:
                label = f"tab:{caption.lower().replace(' ', '_')[:20]}"

            # Generate table
            if format == "latex":
                table_code = self._latex_booktabs(df, caption, label)
            else:
                table_code = self._markdown_table(df, caption)

            # Save
            tables_dir = ws / "analysis" / "tables"
            tables_dir.mkdir(parents=True, exist_ok=True)
            ext = ".tex" if format == "latex" else ".md"
            safe_label = re.sub(r"[^a-zA-Z0-9_]", "_", label)
            output_file = tables_dir / f"{safe_label}{ext}"
            output_file.write_text(table_code, encoding="utf-8")

            return ToolResult(
                success=True,
                content=(
                    f"## Table Generated\n\n"
                    f"**Caption:** {caption}\n"
                    f"**Label:** {label}\n"
                    f"**Dimensions:** {df.shape[0]} rows × {df.shape[1]} columns\n"
                    f"**Saved to:** {output_file.relative_to(ws)}\n\n"
                    f"```{format}\n{table_code}\n```"
                ),
            )
        except Exception as e:
            return ToolResult(success=False, error=f"Table generation failed: {e}")

    @staticmethod
    def _latex_booktabs(df: Any, caption: str, label: str) -> str:
        """Generate LaTeX table with booktabs style."""
        n_cols = len(df.columns)
        col_spec = "l" + "c" * (n_cols - 1) if n_cols > 1 else "l"

        lines: list[str] = [
            r"\begin{table}[htbp]",
            r"  \centering",
            f"  \\caption{{{caption}}}",
            f"  \\label{{{label}}}",
            f"  \\begin{{tabular}}{{{col_spec}}}",
            r"    \toprule",
        ]

        # Header
        header = " & ".join(str(c) for c in df.columns) + r" \\"
        lines.append(f"    {header}")
        lines.append(r"    \midrule")

        # Data rows
        for _, row in df.iterrows():
            cells = []
            for val in row:
                if isinstance(val, float):
                    cells.append(f"{val:.3f}")
                else:
                    cells.append(str(val))
            lines.append(f"    {' & '.join(cells)} \\\\")

        lines.extend(
            [
                r"    \bottomrule",
                r"  \end{tabular}",
                r"\end{table}",
            ]
        )
        return "\n".join(lines)

    @staticmethod
    def _markdown_table(df: Any, caption: str) -> str:
        """Generate Markdown table."""
        lines: list[str] = [f"**Table:** {caption}\n"]
        header = "| " + " | ".join(str(c) for c in df.columns) + " |"
        separator = "|" + "|".join("---" for _ in df.columns) + "|"
        lines.append(header)
        lines.append(separator)
        for _, row in df.iterrows():
            cells = [f"{v:.3f}" if isinstance(v, float) else str(v) for v in row]
            lines.append("| " + " | ".join(cells) + " |")
        return "\n".join(lines)


# ---------------------------------------------------------------------------
# 4. GenerateDiagramTool
# ---------------------------------------------------------------------------


class GenerateDiagramTool(Tool):
    """Generate diagrams using Mermaid or TikZ."""

    def __init__(self, workspace_dir: str = "./workspace") -> None:
        self._workspace_dir = workspace_dir

    @property
    def name(self) -> str:
        return "generate_diagram"

    @property
    def description(self) -> str:
        return (
            "Generate diagrams using Mermaid or TikZ notation. Supports "
            "flowcharts, sequence diagrams, class diagrams, ER diagrams, "
            "conceptual frameworks, and causal DAGs."
        )

    @property
    def parameters(self) -> dict[str, Any]:
        return {
            "type": "object",
            "properties": {
                "diagram_type": {
                    "type": "string",
                    "enum": [
                        "flowchart",
                        "sequence",
                        "class",
                        "er",
                        "conceptual_framework",
                        "causal_dag",
                    ],
                    "description": "Type of diagram to generate",
                },
                "description": {
                    "type": "string",
                    "description": "Natural language description of the diagram",
                },
                "output_path": {
                    "type": "string",
                    "description": "Output file path (optional)",
                },
            },
            "required": ["diagram_type", "description"],
        }

    async def execute(
        self,
        diagram_type: str,
        description: str,
        output_path: str | None = None,
        **kwargs: Any,
    ) -> ToolResult:
        """Generate a diagram from description."""
        try:
            ws = Path(self._workspace_dir)
            diagrams_dir = ws / "analysis" / "diagrams"
            diagrams_dir.mkdir(parents=True, exist_ok=True)

            code = self._generate_mermaid(diagram_type, description)
            ext = ".mmd"

            if output_path:
                out_file = ws / output_path if not Path(output_path).is_absolute() else Path(output_path)
            else:
                safe_name = re.sub(r"[^a-zA-Z0-9_]", "_", diagram_type)
                out_file = diagrams_dir / f"{safe_name}{ext}"

            out_file.parent.mkdir(parents=True, exist_ok=True)
            out_file.write_text(code, encoding="utf-8")

            # Try to render with mermaid-cli
            rendered = self._try_render(out_file)

            render_info = ""
            if rendered:
                render_info = f"\n**Rendered:** {rendered}"
            else:
                render_info = (
                    "\n**Render with:**\n"
                    "  - mmdc -i <file>.mmd -o <file>.svg\n"
                    "  - VS Code: Mermaid Preview extension\n"
                    "  - Online: https://mermaid.live"
                )

            return ToolResult(
                success=True,
                content=(
                    f"## Diagram Generated: {diagram_type}\n\n"
                    f"**Saved to:** {out_file.relative_to(ws)}"
                    f"{render_info}\n\n"
                    f"```mermaid\n{code}\n```"
                ),
            )
        except Exception as e:
            return ToolResult(success=False, error=f"Diagram generation failed: {e}")

    def _generate_mermaid(self, diagram_type: str, description: str) -> str:
        """Generate Mermaid diagram code."""
        generators: dict[str, Any] = {
            "flowchart": self._mermaid_flowchart,
            "sequence": self._mermaid_sequence,
            "class": self._mermaid_class,
            "er": self._mermaid_er,
            "conceptual_framework": self._mermaid_framework,
            "causal_dag": self._mermaid_dag,
        }
        gen = generators.get(diagram_type, self._mermaid_flowchart)
        return gen(description)

    @staticmethod
    def _mermaid_flowchart(description: str) -> str:
        steps = [s.strip() for s in re.split(r"[,;→\->]+", description) if s.strip()]
        if len(steps) < 2:
            steps = description.split()[:5]
        lines = ["graph TD"]
        for i, step in enumerate(steps):
            node_id = chr(65 + i)
            lines.append(f'    {node_id}["{step}"]')
        for i in range(len(steps) - 1):
            lines.append(f"    {chr(65 + i)} --> {chr(65 + i + 1)}")
        return "\n".join(lines)

    @staticmethod
    def _mermaid_sequence(description: str) -> str:
        return (
            f"sequenceDiagram\n"
            f"    %% {description}\n"
            f"    participant A as Actor 1\n"
            f"    participant B as Actor 2\n"
            f"    A->>B: Request\n"
            f"    B-->>A: Response"
        )

    @staticmethod
    def _mermaid_class(description: str) -> str:
        return (
            f"classDiagram\n"
            f"    %% {description}\n"
            f"    class Entity {{\n"
            f"        +attribute1: type\n"
            f"        +method1(): return_type\n"
            f"    }}"
        )

    @staticmethod
    def _mermaid_er(description: str) -> str:
        return (
            f"erDiagram\n"
            f"    %% {description}\n"
            f"    ENTITY1 ||--o{{ ENTITY2 : has\n"
            f"    ENTITY1 {{\n"
            f"        string id PK\n"
            f"        string name\n"
            f"    }}\n"
            f"    ENTITY2 {{\n"
            f"        string id PK\n"
            f"        string entity1_id FK\n"
            f"    }}"
        )

    @staticmethod
    def _mermaid_framework(description: str) -> str:
        return (
            f"graph LR\n"
            f"    %% Conceptual Framework: {description}\n"
            f"    subgraph Independent Variables\n"
            f"        IV1[Variable 1]\n"
            f"        IV2[Variable 2]\n"
            f"    end\n"
            f"    subgraph Mediators\n"
            f"        M1[Mediator]\n"
            f"    end\n"
            f"    subgraph Dependent Variables\n"
            f"        DV1[Outcome]\n"
            f"    end\n"
            f"    IV1 --> M1\n"
            f"    IV2 --> DV1\n"
            f"    M1 --> DV1\n"
            f"    CV[Control Variable] -.-> DV1"
        )

    @staticmethod
    def _mermaid_dag(description: str) -> str:
        return (
            f"graph LR\n"
            f"    %% Causal DAG: {description}\n"
            f"    X[Treatment/Exposure X] --> Y[Outcome Y]\n"
            f"    C[Confounder C] --> X\n"
            f"    C --> Y\n"
            f"    M[Mediator M] --> Y\n"
            f"    X --> M\n"
            f"    U[Unobserved U]:::unobserved -.-> X\n"
            f"    U -.-> Y\n"
            f"    classDef unobserved fill:#f9f,stroke:#333,stroke-dasharray: 5 5"
        )

    @staticmethod
    def _try_render(mmd_path: Path) -> Path | None:
        """Try to render Mermaid to SVG using mmdc."""
        svg_path = mmd_path.with_suffix(".svg")
        try:
            result = subprocess.run(
                ["mmdc", "-i", str(mmd_path), "-o", str(svg_path)],
                capture_output=True,
                text=True,
                timeout=30,
            )
            if result.returncode == 0 and svg_path.exists():
                return svg_path
        except (FileNotFoundError, subprocess.TimeoutExpired):
            pass
        return None


# ---------------------------------------------------------------------------
# 5. ConvertFigureTikzTool
# ---------------------------------------------------------------------------


class ConvertFigureTikzTool(Tool):
    """Convert matplotlib figure to TikZ/PGFplots for LaTeX."""

    def __init__(self, workspace_dir: str = "./workspace") -> None:
        self._workspace_dir = workspace_dir

    @property
    def name(self) -> str:
        return "convert_figure_tikz"

    @property
    def description(self) -> str:
        return (
            "Convert a matplotlib figure (.py script or image) to TikZ/PGFplots "
            "code for native LaTeX rendering. Uses tikzplotlib if available, "
            "otherwise generates an \\includegraphics wrapper."
        )

    @property
    def parameters(self) -> dict[str, Any]:
        return {
            "type": "object",
            "properties": {
                "figure_path": {
                    "type": "string",
                    "description": "Path to .py script or image file",
                },
                "output_path": {
                    "type": "string",
                    "description": "Output path for TikZ .tex file (optional)",
                },
            },
            "required": ["figure_path"],
        }

    async def execute(
        self,
        figure_path: str,
        output_path: str | None = None,
        **kwargs: Any,
    ) -> ToolResult:
        """Convert figure to TikZ."""
        try:
            ws = Path(self._workspace_dir)
            fig_path = ws / figure_path if not Path(figure_path).is_absolute() else Path(figure_path)

            if not fig_path.exists():
                return ToolResult(success=False, error=f"Figure file not found: {fig_path}")

            if output_path:
                out_path = ws / output_path if not Path(output_path).is_absolute() else Path(output_path)
            else:
                out_path = fig_path.with_suffix(".tex")

            out_path.parent.mkdir(parents=True, exist_ok=True)

            if fig_path.suffix == ".py":
                tikz_code = self._convert_script(fig_path)
            elif fig_path.suffix in (".png", ".pdf", ".jpg", ".jpeg", ".svg"):
                tikz_code = self._convert_image(fig_path)
            else:
                return ToolResult(
                    success=False,
                    error=f"Unsupported file type: {fig_path.suffix}",
                )

            out_path.write_text(tikz_code, encoding="utf-8")

            return ToolResult(
                success=True,
                content=(
                    f"## Figure Converted to TikZ\n\n"
                    f"**Input:** {fig_path.name}\n"
                    f"**Output:** {out_path}\n\n"
                    f"```latex\n{tikz_code}\n```"
                ),
            )
        except Exception as e:
            return ToolResult(success=False, error=f"Figure conversion failed: {e}")

    def _convert_script(self, script_path: Path) -> str:
        """Try tikzplotlib on a matplotlib script."""
        try:
            import matplotlib
            import tikzplotlib  # type: ignore[import-untyped]

            matplotlib.use("Agg")
            import matplotlib.pyplot as plt

            namespace: dict[str, Any] = {}
            exec(script_path.read_text(encoding="utf-8"), namespace)
            fig = plt.gcf()
            tikz_code = tikzplotlib.get_tikz_code(figure=fig)
            plt.close("all")
            return tikz_code
        except ImportError:
            return self._includegraphics_wrapper(script_path.with_suffix(".pdf"))

    def _convert_image(self, image_path: Path) -> str:
        """Try tikzplotlib on an image, fall back to includegraphics."""
        try:
            import matplotlib.image as mpimg
            import matplotlib.pyplot as plt
            import tikzplotlib  # type: ignore[import-untyped]

            img = mpimg.imread(str(image_path))
            fig, ax = plt.subplots()
            ax.imshow(img)
            ax.axis("off")
            tikz_code = tikzplotlib.get_tikz_code(figure=fig)
            plt.close(fig)
            return tikz_code
        except (ImportError, Exception):
            return self._includegraphics_wrapper(image_path)

    @staticmethod
    def _includegraphics_wrapper(image_path: Path) -> str:
        """Generate LaTeX includegraphics wrapper."""
        return textwrap.dedent(f"""\
            % Auto-generated includegraphics wrapper
            % For native TikZ, install tikzplotlib: pip install tikzplotlib
            \\begin{{figure}}[htbp]
                \\centering
                \\includegraphics[width=0.8\\textwidth]{{{image_path.name}}}
                \\caption{{TODO: Add caption}}
                \\label{{fig:{image_path.stem}}}
            \\end{{figure}}""")
